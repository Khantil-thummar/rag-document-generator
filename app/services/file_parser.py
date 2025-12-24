"""
File parser service for extracting text from different file formats.
Supports: .txt, .pdf, .docx
"""

import io
from pypdf import PdfReader
from docx import Document


# Supported file extensions
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def get_file_extension(filename: str) -> str:
    """Get the lowercase file extension from filename."""
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def is_supported_file(filename: str) -> bool:
    """Check if the file type is supported."""
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


def extract_text_from_txt(content: bytes) -> str:
    """Extract text from a .txt file."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        # Try other common encodings
        for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file with supported encodings")


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from a .pdf file."""
    try:
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Could not parse PDF file: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from a .docx file."""
    try:
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Could not parse DOCX file: {str(e)}")


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract text from a file based on its extension.
    
    Args:
        filename: Name of the file (used to determine type)
        content: Raw file content as bytes
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is unsupported or parsing fails
    """
    extension = get_file_extension(filename)
    
    if extension == ".txt":
        return extract_text_from_txt(content)
    elif extension == ".pdf":
        return extract_text_from_pdf(content)
    elif extension == ".docx":
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")

